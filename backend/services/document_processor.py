"""
Document processing service for extracting text from various file formats
"""
import io
import os
import asyncio
from typing import Optional, Dict, Any, BinaryIO, List
from pathlib import Path
import tempfile

import PyPDF2
import pdfplumber
from docx import Document
from openpyxl import load_workbook
from PIL import Image
import pytesseract

from loguru import logger
from config.settings import settings
from services.claude_service import openrouter_service, specification_validator
from services.vector_search_service import vector_search_service
from models.database import DocumentChunk
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select


class DocumentProcessor:
    """Process various document formats and extract text"""
    
    def __init__(self):
        self.ocr_language = settings.OCR_LANGUAGE
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
        self.max_chunks = 1000000  # Allow very large documents (increased limit for big PDFs)
    
    async def process_file(
        self,
        file_content: bytes,
        filename: str,
        file_type: str,
        document_id: Optional[int] = None,
        db: Optional[AsyncSession] = None
    ) -> Dict[str, Any]:
        """
        Process uploaded file and extract text with enhanced chunking and embedding

        Args:
            file_content: File content as bytes
            filename: Original filename
            file_type: File type/extension
            document_id: Optional document ID for storing chunks
            db: Optional database session for storing chunks

        Returns:
            Dictionary with extracted text and metadata
        """
        file_type = file_type.lower().lstrip('.')

        logger.info(f"Processing file {filename} (type: {file_type}) - enhanced processing pipeline")

        # Extract text based on file type
        if file_type == 'pdf':
            result = await self._process_pdf(file_content, filename)
        elif file_type == 'docx':
            result = await self._process_docx(file_content, filename)
        elif file_type == 'xlsx':
            result = await self._process_xlsx(file_content, filename)
        elif file_type in ['txt']:
            result = await self._process_txt(file_content, filename)
        elif file_type in ['png', 'jpg', 'jpeg']:
            result = await self._process_image(file_content, filename)
        else:
            # Fallback for unsupported types
            result = {
                'text': f"Unsupported file type: {file_type}",
                'metadata': {'filename': filename, 'file_type': file_type, 'size': len(file_content)},
                'chunks': []
            }

        # Generate embeddings for chunks if document_id and db are provided
        if document_id and db and result.get('chunks'):
            await self._generate_and_store_embeddings(result['chunks'], document_id, db)

        return result
        
    
    async def _process_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from PDF"""
        text = ""
        metadata = {}

        try:
            logger.info(f"Processing PDF {filename}: Starting pdfplumber text extraction")
            # First, try pdfplumber for better text extraction
            loop = asyncio.get_event_loop()
            pdf_result = await loop.run_in_executor(None, self._extract_pdf_text_sync, file_content)
            text = pdf_result['text']
            metadata.update(pdf_result['metadata'])
            logger.info(f"Processing PDF {filename}: pdfplumber text extraction complete")

            # If no text extracted, try OCR
            if len(text.strip()) < 100:
                logger.info(f"PDF {filename} has little text, trying OCR")
                text = await self._ocr_pdf(file_content)
                metadata['ocr_used'] = True
            else:
                metadata['ocr_used'] = False

            logger.info(f"Processing PDF {filename}: Starting metadata extraction")
            # Extract metadata
            metadata_result = await loop.run_in_executor(None, self._extract_pdf_metadata_sync, file_content)
            metadata.update(metadata_result)
            logger.info(f"Processing PDF {filename}: Metadata extraction complete")

        except Exception as e:
            logger.error(f"Error processing PDF {filename}: {e}")
            raise

        return {
            'text': text,
            'metadata': metadata,
            'chunks': self._chunk_text(text)
        }

    async def extract_text_sample(self, file_content: bytes, filename: str, file_type: str) -> str:
        """Extract a sample of text for validation purposes"""
        try:
            if file_type.lower() == 'pdf':
                loop = asyncio.get_event_loop()
                pdf_result = await loop.run_in_executor(None, self._extract_pdf_text_sync, file_content)
                return pdf_result['text'][:500]  # Return first 500 characters
            elif file_type.lower() == 'docx':
                loop = asyncio.get_event_loop()
                result = await loop.run_in_executor(None, self._process_docx_sync, file_content, filename)
                return result['text'][:500]
            elif file_type.lower() == 'txt':
                return file_content.decode('utf-8')[:500]
            else:
                return "Unsupported file type for text extraction"
        except Exception as e:
            logger.error(f"Error extracting text sample from {filename}: {e}")
            return ""
    
    async def _ocr_pdf(self, file_content: bytes) -> str:
        """Perform OCR on PDF pages"""
        loop = asyncio.get_event_loop()
        text = await loop.run_in_executor(None, self._ocr_pdf_sync, file_content)
        return text

    def _ocr_pdf_sync(self, file_content: bytes) -> str:
        """Perform OCR on PDF pages (synchronous)"""
        text = ""

        try:
            # Convert PDF to images and OCR
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
                tmp_file.write(file_content)
                tmp_path = tmp_file.name

            try:
                # Use pdfplumber to extract images
                with pdfplumber.open(tmp_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        # Convert page to image
                        img = page.to_image(resolution=300)

                        # Perform OCR
                        page_text = pytesseract.image_to_string(
                            img.original,
                            lang=self.ocr_language
                        )
                        text += f"\n\n--- Page {page_num + 1} ---\n\n{page_text}"
            finally:
                os.unlink(tmp_path)

        except Exception as e:
            logger.error(f"Error during OCR: {e}")
            raise

        return text
    def _extract_pdf_text_sync(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text from PDF synchronously"""
        text = ""
        metadata = {}

        with io.BytesIO(file_content) as pdf_file:
            with pdfplumber.open(pdf_file) as pdf:
                metadata['pages'] = len(pdf.pages)

                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"

        return {'text': text, 'metadata': metadata}

    def _extract_pdf_metadata_sync(self, file_content: bytes) -> Dict[str, Any]:
        """Extract metadata from PDF synchronously"""
        metadata = {}

        with io.BytesIO(file_content) as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            if pdf_reader.metadata:
                metadata['title'] = pdf_reader.metadata.get('/Title', '')
                metadata['author'] = pdf_reader.metadata.get('/Author', '')
                metadata['subject'] = pdf_reader.metadata.get('/Subject', '')

        return metadata
    
    async def _process_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from DOCX"""
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(None, self._process_docx_sync, file_content, filename)
        return result

    def _process_docx_sync(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from DOCX (synchronous)"""
        text = ""
        metadata = {}

        try:
            with io.BytesIO(file_content) as docx_file:
                doc = Document(docx_file)

                # Extract text from paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        text += para.text + "\n"

                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row.cells)
                        text += row_text + "\n"

                # Extract metadata
                core_props = doc.core_properties
                metadata['title'] = core_props.title or ""
                metadata['author'] = core_props.author or ""
                metadata['subject'] = core_props.subject or ""

        except Exception as e:
            logger.error(f"Error processing DOCX {filename}: {e}")
            raise

        return {
            'text': text,
            'metadata': metadata,
            'chunks': self._chunk_text(text)
        }
    
    async def _process_xlsx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from XLSX"""
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(None, self._process_xlsx_sync, file_content, filename)
        return result

    def _process_xlsx_sync(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from XLSX (synchronous)"""
        text = ""
        metadata = {}

        try:
            with io.BytesIO(file_content) as xlsx_file:
                workbook = load_workbook(xlsx_file, data_only=True)

                metadata['sheets'] = workbook.sheetnames

                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text += f"\n\n=== Sheet: {sheet_name} ===\n\n"

                    for row in sheet.iter_rows(values_only=True):
                        row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                        if row_text.strip():
                            text += row_text + "\n"

        except Exception as e:
            logger.error(f"Error processing XLSX {filename}: {e}")
            raise

        return {
            'text': text,
            'metadata': metadata,
            'chunks': self._chunk_text(text)
        }
    
    async def _process_txt(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from TXT"""
        try:
            text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-2']:
                try:
                    text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError(f"Unable to decode text file {filename}")
        
        return {
            'text': text,
            'metadata': {},
            'chunks': self._chunk_text(text)
        }
    
    async def _process_image(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from image using OCR"""
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(None, self._process_image_sync, file_content, filename)
        return result

    def _process_image_sync(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from image using OCR (synchronous)"""
        text = ""
        metadata = {}

        try:
            with io.BytesIO(file_content) as img_file:
                image = Image.open(img_file)

                metadata['size'] = image.size
                metadata['format'] = image.format

                # Perform OCR
                text = pytesseract.image_to_string(image, lang=self.ocr_language)

        except Exception as e:
            logger.error(f"Error processing image {filename}: {e}")
            raise

        return {
            'text': text,
            'metadata': metadata,
            'chunks': self._chunk_text(text)
        }
    
    def _chunk_text(self, text: str) -> list[str]:
        """Split text into chunks for embedding"""
        if not text:
            return []

        chunks = []
        start = 0
        text_len = len(text)

        while start < text_len:
            end = start + self.chunk_size

            # Try to break at sentence or paragraph boundary
            if end < text_len:
                # Look for paragraph break
                break_pos = text.rfind('\n\n', start, end)
                if break_pos == -1:
                    # Look for sentence break
                    break_pos = text.rfind('. ', start, end)
                if break_pos == -1:
                    # Look for any space
                    break_pos = text.rfind(' ', start, end)
                if break_pos != -1 and break_pos > start:
                    end = break_pos + 1

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - self.chunk_overlap

            # Prevent infinite loop by ensuring progress
            if start >= text_len:
                break
            if len(chunks) > self.max_chunks:  # Use configurable limit
                logger.warning(f"Too many chunks ({len(chunks)}), limit is {self.max_chunks}, truncating text. Document size: ~{len(text)} chars, chunk size: {self.chunk_size}, overlap: {self.chunk_overlap}")
                break

        return chunks

    async def _generate_and_store_embeddings(
        self,
        chunks: List[str],
        document_id: int,
        db: AsyncSession,
        use_batch_processing: bool = True,
        batch_size: int = 10
    ) -> None:
        """
        Generate embeddings for text chunks and store them in the database

        Args:
            chunks: List of text chunks
            document_id: Document ID to associate chunks with
            db: Database session
            use_batch_processing: Whether to use batch processing for embeddings
            batch_size: Size of batches for processing
        """
        try:
            logger.info(f"Generating embeddings for {len(chunks)} chunks for document {document_id}")

            if use_batch_processing and len(chunks) > 1:
                # Use batch processing for better performance
                logger.info(f"Using batch processing with batch size {batch_size}")
                embeddings = await vector_search_service.batch_generate_embeddings(chunks, batch_size)

                for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                    await self._store_single_chunk(db, document_id, chunk, embedding, i)

            else:
                # Process chunks individually (fallback or for single chunks)
                for i, chunk in enumerate(chunks):
                    embedding = await openrouter_service.generate_embedding(chunk)
                    await self._store_single_chunk(db, document_id, chunk, embedding, i)

            await db.commit()
            logger.info(f"Successfully stored {len(chunks)} chunks with embeddings for document {document_id}")

        except Exception as e:
            logger.error(f"Error generating and storing embeddings for document {document_id}: {e}")
            await db.rollback()
            raise

    async def _store_single_chunk(
        self,
        db: AsyncSession,
        document_id: int,
        chunk: str,
        embedding: List[float],
        chunk_index: int
    ) -> None:
        """
        Store a single chunk with its embedding

        Args:
            db: Database session
            document_id: Document ID
            chunk: Text chunk
            embedding: Embedding vector
            chunk_index: Index of the chunk
        """
        chunk_record = DocumentChunk(
            document_id=document_id,
            content=chunk,
            chunk_index=chunk_index,
            chunk_type="paragraph",  # Default type, could be enhanced to detect headers/tables
            embedding=embedding,
            chunk_metadata={
                "chunk_length": len(chunk),
                "start_position": chunk_index * (self.chunk_size - self.chunk_overlap)
            }
        )

        db.add(chunk_record)
        logger.info(f"Stored chunk {chunk_index+1} for document {document_id}")


# Singleton instance
document_processor = DocumentProcessor()
