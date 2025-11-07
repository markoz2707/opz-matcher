"""
API routes for OPZ creation mode - generating OPZ documents
"""
from fastapi import APIRouter, HTTPException, Depends, BackgroundTasks
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from typing import List, Optional, Dict, Any
from pydantic import BaseModel
from datetime import datetime
import os
import tempfile

from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from services.database import get_db
from services.claude_service import openrouter_service
from services.storage_service import storage_service
from models.database import OPZDocument, DeviceCategory
from api.dependencies import get_current_user
from loguru import logger


router = APIRouter()


class OPZCreateRequest(BaseModel):
    title: str
    category: DeviceCategory
    configuration: Dict[str, Any]
    selected_vendors: List[str]
    template_type: str = "standard"


class OPZRefineRequest(BaseModel):
    feedback: str


class OPZResponse(BaseModel):
    opz_id: int
    title: str
    category: str
    status: str
    generated_text: Optional[str] = None
    created_at: datetime
    updated_at: Optional[datetime] = None


@router.post("/create", response_model=OPZResponse)
async def create_opz(
    request: OPZCreateRequest,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """
    Create a new OPZ document
    
    This endpoint:
    1. Creates an OPZ record in the database
    2. Generates OPZ text using Claude in the background
    3. Returns the OPZ ID for tracking
    """
    
    # Create OPZ record
    opz_doc = OPZDocument(
        owner_id=current_user.id,
        title=request.title,
        category=request.category,
        requirements=request.configuration,
        selected_vendors=request.selected_vendors,
        status="generating"
    )
    
    db.add(opz_doc)
    await db.commit()
    await db.refresh(opz_doc)
    
    # Generate OPZ in background
    background_tasks.add_task(
        generate_opz_task,
        opz_doc.id,
        request.category.value,
        request.configuration,
        request.selected_vendors,
        request.template_type
    )
    
    return OPZResponse(
        opz_id=opz_doc.id,
        title=opz_doc.title,
        category=opz_doc.category.value,
        status=opz_doc.status,
        created_at=opz_doc.created_at,
        updated_at=opz_doc.updated_at
    )


async def generate_opz_task(
    opz_id: int,
    category: str,
    configuration: Dict[str, Any],
    vendors: List[str],
    template_type: str
):
    """Background task to generate OPZ text"""
    from services.database import AsyncSessionLocal
    from sqlalchemy import select
    
    async with AsyncSessionLocal() as db:
        try:
            # Generate OPZ text using OpenRouter
            opz_text = await openrouter_service.generate_opz(
                category=category,
                configuration=configuration,
                vendors=vendors,
                template_type=template_type
            )
            
            # Update OPZ record
            result = await db.execute(
                select(OPZDocument).where(OPZDocument.id == opz_id)
            )
            opz_doc = result.scalar_one()
            
            opz_doc.generated_text = opz_text
            opz_doc.status = "generated"
            opz_doc.generated_at = datetime.utcnow()
            
            await db.commit()
            
        except Exception as e:
            # Update with error status
            result = await db.execute(
                select(OPZDocument).where(OPZDocument.id == opz_id)
            )
            opz_doc = result.scalar_one()
            
            opz_doc.status = "error"
            await db.commit()
            
            logger.error(f"Error generating OPZ {opz_id}: {e}")


@router.get("/{opz_id}", response_model=OPZResponse)
async def get_opz(
    opz_id: int,
    db: AsyncSession = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """Get OPZ document details"""
    from sqlalchemy import select
    
    result = await db.execute(
        select(OPZDocument).where(OPZDocument.id == opz_id)
    )
    opz_doc = result.scalar_one_or_none()
    
    if not opz_doc:
        raise HTTPException(status_code=404, detail="OPZ document not found")
    
    # Check ownership
    if opz_doc.owner_id != current_user.id and not current_user.is_superuser:
        raise HTTPException(status_code=403, detail="Access denied")
    
    return OPZResponse(
        opz_id=opz_doc.id,
        title=opz_doc.title,
        category=opz_doc.category.value,
        status=opz_doc.status,
        generated_text=opz_doc.generated_text,
        created_at=opz_doc.created_at,
        updated_at=opz_doc.updated_at
    )


@router.post("/{opz_id}/refine", response_model=OPZResponse)
async def refine_opz(
    opz_id: int,
    request: OPZRefineRequest,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """
    Refine OPZ based on user feedback
    """
    from sqlalchemy import select
    
    result = await db.execute(
        select(OPZDocument).where(OPZDocument.id == opz_id)
    )
    opz_doc = result.scalar_one_or_none()
    
    if not opz_doc:
        raise HTTPException(status_code=404, detail="OPZ document not found")
    
    if opz_doc.owner_id != current_user.id and not current_user.is_superuser:
        raise HTTPException(status_code=403, detail="Access denied")
    
    if not opz_doc.generated_text:
        raise HTTPException(status_code=400, detail="OPZ not yet generated")
    
    # Update status
    opz_doc.status = "refining"
    await db.commit()
    
    # Refine in background
    background_tasks.add_task(
        refine_opz_task,
        opz_id,
        opz_doc.generated_text,
        request.feedback
    )
    
    return OPZResponse(
        opz_id=opz_doc.id,
        title=opz_doc.title,
        category=opz_doc.category.value,
        status=opz_doc.status,
        created_at=opz_doc.created_at,
        updated_at=opz_doc.updated_at
    )


async def refine_opz_task(
    opz_id: int,
    current_text: str,
    feedback: str
):
    """Background task to refine OPZ"""
    from services.database import AsyncSessionLocal
    from sqlalchemy import select
    
    async with AsyncSessionLocal() as db:
        try:
            # Refine OPZ using OpenRouter
            refined_text = await openrouter_service.refine_opz(
                current_opz=current_text,
                feedback=feedback
            )
            
            # Update OPZ record
            result = await db.execute(
                select(OPZDocument).where(OPZDocument.id == opz_id)
            )
            opz_doc = result.scalar_one()
            
            opz_doc.generated_text = refined_text
            opz_doc.status = "generated"
            opz_doc.updated_at = datetime.utcnow()
            
            await db.commit()
            
        except Exception as e:
            logger.error(f"Error refining OPZ {opz_id}: {e}")


@router.get("/{opz_id}/download")
async def download_opz(
    opz_id: int,
    db: AsyncSession = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """
    Download OPZ as DOCX file
    """
    from sqlalchemy import select
    
    result = await db.execute(
        select(OPZDocument).where(OPZDocument.id == opz_id)
    )
    opz_doc = result.scalar_one_or_none()
    
    if not opz_doc:
        raise HTTPException(status_code=404, detail="OPZ document not found")
    
    if opz_doc.owner_id != current_user.id and not current_user.is_superuser:
        raise HTTPException(status_code=403, detail="Access denied")
    
    if not opz_doc.generated_text:
        raise HTTPException(status_code=400, detail="OPZ not yet generated")
    
    # Generate DOCX file
    docx_path = await create_opz_docx(
        opz_doc.title,
        opz_doc.generated_text,
        opz_doc.category.value
    )
    
    # Update status
    opz_doc.status = "downloaded"
    await db.commit()
    
    return FileResponse(
        docx_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"OPZ_{opz_doc.title.replace(' ', '_')}_{opz_doc.id}.docx"
    )


async def create_opz_docx(
    title: str,
    content: str,
    category: str
) -> str:
    """
    Create a formatted DOCX file from OPZ text
    """
    
    # Create document
    doc = DocxDocument()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Add title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_paragraph(f"Kategoria: {category}")
    doc.add_paragraph(f"Data utworzenia: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph()  # Empty line
    
    # Parse and format content
    # Split by sections (assuming numbered sections)
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        
        # Detect headers (lines ending with :, or starting with numbers)
        if line.endswith(':') or (len(line) > 0 and line[0].isdigit() and '. ' in line[:10]):
            # This is likely a header
            doc.add_heading(line, level=1)
        elif line.startswith('- ') or line.startswith('• '):
            # This is a bullet point
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            # Regular paragraph
            doc.add_paragraph(line)
    
    # Add footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Dokument wygenerowany przez OPZ Product Matcher"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        doc.save(tmp.name)
        return tmp.name


@router.get("/list", response_model=List[OPZResponse])
async def list_user_opz_documents(
    db: AsyncSession = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """List all OPZ documents for the current user"""
    from sqlalchemy import select
    
    query = select(OPZDocument).where(OPZDocument.owner_id == current_user.id)
    query = query.order_by(OPZDocument.created_at.desc())
    
    result = await db.execute(query)
    opz_docs = result.scalars().all()
    
    return [
        OPZResponse(
            opz_id=doc.id,
            title=doc.title,
            category=doc.category.value,
            status=doc.status,
            created_at=doc.created_at,
            updated_at=doc.updated_at
        )
        for doc in opz_docs
    ]


@router.delete("/{opz_id}", response_model=dict)
async def delete_opz(
    opz_id: int,
    db: AsyncSession = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """Delete an OPZ document"""
    from sqlalchemy import select, delete

    logger.info(f"DEBUG: Starting OPZ deletion for ID {opz_id} by user: {current_user.username}")

    result = await db.execute(
        select(OPZDocument).where(OPZDocument.id == opz_id)
    )
    opz_doc = result.scalar_one_or_none()

    if not opz_doc:
        logger.warning(f"DEBUG: OPZ deletion failed: OPZ {opz_id} not found")
        raise HTTPException(status_code=404, detail="OPZ document not found")

    logger.info(f"DEBUG: Found OPZ {opz_id}, title: {opz_doc.title}, owner: {opz_doc.owner_id}")

    if opz_doc.owner_id != current_user.id and not current_user.is_superuser:
        logger.warning(f"DEBUG: Access denied for OPZ {opz_id} - user {current_user.id} vs owner {opz_doc.owner_id}")
        raise HTTPException(status_code=403, detail="Access denied")

    logger.info(f"DEBUG: Deleting OPZ record {opz_id} from database")
    await db.execute(
        delete(OPZDocument).where(OPZDocument.id == opz_id)
    )
    await db.commit()
    logger.info(f"DEBUG: OPZ {opz_id} deletion completed successfully")

    return {"message": "OPZ document deleted successfully"}
